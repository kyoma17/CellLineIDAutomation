import tkinter as tk
from tkinter import ttk
import docx
import pandas
import win32com.client
import os
import psutil

def fillTemplate(replacementsDictionary):
    # Helper function to fill the template with the data from the dictionary
    sampleName = replacementsDictionary["_SAMPLE_NAME"]

    close_word_processes()

    # Determine which template to use based on the test name
    if replacementsDictionary["test"].values[0] == "GenePrint_24_POP7_Panels_v1.0":
        document = docx.Document('CellLineTemplateGP24.docx')
    elif replacementsDictionary["test"].values[0] == "GenePrint_10_v1.1":
        document = docx.Document('CellLineTemplateGP10.docx')

    # Iterate through the tables in the document and replace the text from the dictionary, Preserve the font attributes
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:

                    for run in paragraph.runs:

                        for key in replacementsDictionary.keys():
                            if key in run.text:
                                # If the key is any of the special keys, peform a special action to color the text

                                # Flag to color the text red
                                keys_to_check = ["D5S818_bM", "D13S317_bM", "D7S820_bM", "D16S539_bM",
                                                 "vWA_bM", "TH01_bM", "AMEL_bM", "TPOX_bM", "CSF1PO_bM", "D21S11_bM"]

                                if key in keys_to_check:

                                    # run color replacement
                                    NumbersToColorRed = []
                                    # Get the value of the key and seperate by comma
                                    dictionaryValue = replacementsDictionary[key].split(",")

                                    textToReplaceWith = replacementsDictionary[key]

                                    alleleLookupValue1 = replacementsDictionary[key[:-3] + "_1"]
                                    alleleLookupValue2 = replacementsDictionary[key[:-3] + "_2"]
                                    alleleLookupValue3 = replacementsDictionary[key[:-3] + "_3"]
                                    alleleLookupValue4 = replacementsDictionary[key[:-3] + "_4"]

                                    ValuesToCheck = []

                                    for allele in [alleleLookupValue1, alleleLookupValue2, alleleLookupValue3, alleleLookupValue4]:
                                        # if nan skip
                                        if not pandas.isna(allele):
                                            ValuesToCheck.append(allele)

                                    for value in dictionaryValue:
                                        if value in ValuesToCheck:
                                            NumbersToColorRed.append(value)

                                # run Normal replacement
                                # Save the current font attributes
                                font_name = run.font.name
                                font_size = run.font.size
                                font_bold = run.font.bold
                                font_italic = run.font.italic

                                # Replace the text
                                if replacementsDictionary[key] is None or pandas.isna(replacementsDictionary[key]):
                                    run.text = run.text.replace(key, "")
                                else:
                                    run.text = run.text.replace(
                                        key, str(replacementsDictionary[key]))

                                # Apply the saved font attributes to the Font object
                                run.font.name = font_name
                                run.font.size = font_size
                                run.font.bold = font_bold
                                run.font.italic = font_italic

    # Second pass: Iterate through the tables in the document and replace the text from the dictionary
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in replacementsDictionary.keys():
                        if key in paragraph.text:
                            # Enter the replacement text in the document, if the value is None or nan, then replace with an empty string
                            if replacementsDictionary[key] is None or pandas.isna(replacementsDictionary[key]):
                                paragraph.text = paragraph.text.replace(
                                    key, "")
                            else:
                                paragraph.text = paragraph.text.replace(
                                    key, str(replacementsDictionary[key]))

    # Inject and run VBA code to the document

    # Save the modified document to the CellLineTEMP folder in the current directory
    document.save('CellLineTEMP/' + sampleName + '.docx')

    injectAndRunRedCodeVBA('CellLineTEMP/' + sampleName + '.docx')

    print("Done with " + sampleName)
    # Print line empty line
    print("")

########################################################################################################################
def injectAndRunRedCodeVBA(fileName):
    # Inject and run VBA code to the document

    # Current Directory
    currentDirectory = os.getcwd()

    filePath = os.path.join(currentDirectory, fileName)

    # Open the document with win32com turn on the visible mode
    word = win32com.client.DispatchEx("Word.Application")

    # Open the document in read only mode
    doc = word.Documents.Open(filePath, ReadOnly=True)

    # Load the VBA code from the file
    with open("MSO/RedHighlighter.bas", "r") as f:
        vbaCode = f.read()

    # Inject the VBA code to the document
    doc.VBProject.VBComponents.Add(1).CodeModule.AddFromString(vbaCode)

    # Run the VBA code
    word.Run("ChangeMatchingToRed")

    # Save the document 
    doc.Save()

    # Close the document
    doc.Close()

    # Quit the word application
    word.Quit()

########################################################################################################################
def close_word_processes():
    for process in psutil.process_iter(['pid', 'name']):
        if process.info['name'] == 'WINWORD.EXE':  # Check for the Word process name
            try:
                process.kill()  # Terminate the Word process
                print(f"Closed Word process with PID: {process.info['pid']}")
            except psutil.AccessDenied:
                print(f"Access denied to terminate Word process with PID: {process.info['pid']}")


def close_microsoft_word():
    # close word application if it is open
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Quit()
    except:
        pass