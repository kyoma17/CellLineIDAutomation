# Python-Docx Template Generator to create the formatted Word document
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

light_grey = "D3D3D3"

def create_template():
    # Initialize a Document object
    doc = Document()

    table = doc.add_table(rows=2, cols=5)

    for cell in table.rows[0].cells:
        cell.width = Inches(4.0)

    table.style = 'Table Grid'

    a = table.cell(1, 0)
    b = table.cell(1, 1)
    A = a.merge(b)

    # Row 1 of first table
    table.cell(0, 0).text = '_SAMPLE_NUMBER'
    set_cell_font_bold(table.cell(0, 0))


    table.cell(0, 1).text = 'Sample ID:'
    set_cell_background(table.cell(0, 1), light_grey)
    
    table.cell(0, 2).text = '_SAMPLE_NAME'
    table.cell(0, 3).text = 'Dataset Best Match Cell Name:'
    set_cell_background(table.cell(0, 3), light_grey)

    table.cell(0, 4).text = '_bMatchName'

    # Row 2 of first table
    table.cell(1, 0).text = 'Database Best Match Score:'
    set_cell_background(table.cell(1, 0), light_grey)

    table.cell(1, 2).text = '_bMatchScore'
    table.cell(1, 3).text = 'Dataset Best Match Cell Line No:'
    set_cell_background(table.cell(1, 3), light_grey)

    table.cell(1, 4).text = '_bMatchCellLineNo'

    doc.add_paragraph('')

    table2 = doc.add_table(rows=14, cols=6)

    a = table2.cell(0, 1)
    b = table2.cell(0, 6)
    A = a.merge(b)

    table2.style = 'Table Grid'

    table2.cell(0, 0).text = 'Methodology:'
    set_cell_font_bold(table2.cell(0, 0))

    table2.cell(0, 1).text = 'Authenticity for all cell lines used in this study was performed using the Promega GenePrint 10 System and following the protocol described in ANSI/ATCC ASN-0002-2011.  The STR alleles were searched on the _dataset Database.'

    a = table2.cell(2, 0)
    b = table2.cell(2, 6)
    A = a.merge(b)

    table2.cell(2, 0).text = 'Marker Characterization'

    table2.cell(3, 0).text = 'Marker'
    set_cell_background(table2.cell(3, 0), light_grey)

    table2.cell(3, 1).text = 'Allele 1'
    set_cell_background(table2.cell(3, 1), light_grey)

    table2.cell(3, 2).text = 'Allele 2'
    set_cell_background(table2.cell(3, 2), light_grey)

    table2.cell(3, 3).text = 'Allele 3'
    set_cell_background(table2.cell(3, 3), light_grey)

    table2.cell(3, 4).text = 'Allele 4'
    set_cell_background(table2.cell(3, 4), light_grey)

    table2.cell(3, 5).text = 'Dataset Best Match Profile'

    table2.cell(4, 0).text = 'D5S818'
    table2.cell(4, 1).text = 'D5S818_1'
    table2.cell(4, 2).text = 'D5S818_2'
    table2.cell(4, 3).text = 'D5S818_3'
    table2.cell(4, 4).text = 'D5S818_4'
    table2.cell(4, 5).text = 'D5S818_bM'

    table2.cell(5, 0).text = 'D13S317'
    table2.cell(5, 1).text = 'D13S317_1'
    table2.cell(5, 2).text = 'D13S317_2'
    table2.cell(5, 3).text = 'D13S317_3'
    table2.cell(5, 4).text = 'D13S317_4'
    table2.cell(5, 5).text = 'D13S317_bM'

    table2.cell(6, 0).text = 'D7S820'
    table2.cell(6, 1).text = 'D7S820_1'
    table2.cell(6, 2).text = 'D7S820_2'
    table2.cell(6, 3).text = 'D7S820_3'
    table2.cell(6, 4).text = 'D7S820_4'
    table2.cell(6, 5).text = 'D7S820_bM'

    table2.cell(7, 0).text = 'D16S539'
    table2.cell(7, 1).text = 'D16S539_1'
    table2.cell(7, 2).text = 'D16S539_2'
    table2.cell(7, 3).text = 'D16S539_3'
    table2.cell(7, 4).text = 'D16S539_4'
    table2.cell(7, 5).text = 'D16S539_bM'

    table2.cell(8, 0).text = 'vWA'
    table2.cell(8, 1).text = 'vWA_1'
    table2.cell(8, 2).text = 'vWA_2'
    table2.cell(8, 3).text = 'vWA_3'
    table2.cell(8, 4).text = 'vWA_4'
    table2.cell(8, 5).text = 'vWA_bM'

    table2.cell(9, 0).text = 'TH01'
    table2.cell(9, 1).text = 'TH01_1'
    table2.cell(9, 2).text = 'TH01_2'
    table2.cell(9, 3).text = 'TH01_3'
    table2.cell(9, 4).text = 'TH01_4'
    table2.cell(9, 5).text = 'TH01_bM'

    table2.cell(10, 0).text = 'AMEL'
    table2.cell(10, 1).text = 'AMEL_1'
    table2.cell(10, 2).text = 'AMEL_2'
    table2.cell(10, 3).text = 'AMEL_3'
    table2.cell(10, 4).text = 'AMEL_4'
    table2.cell(10, 5).text = 'AMEL_bM'

    table2.cell(11, 0).text = 'TPOX'
    table2.cell(11, 1).text = 'TPOX_1'
    table2.cell(11, 2).text = 'TPOX_2'
    table2.cell(11, 3).text = 'TPOX_3'
    table2.cell(11, 4).text = 'TPOX_4'
    table2.cell(11, 5).text = 'TPOX_bM'

    table2.cell(12, 0).text = 'CSF1PO'
    table2.cell(12, 1).text = 'CSF1PO_1'
    table2.cell(12, 2).text = 'CSF1PO_2'
    table2.cell(12, 3).text = 'CSF1PO_3'
    table2.cell(12, 4).text = 'CSF1PO_4'
    table2.cell(12, 5).text = 'CSF1PO_bM'

    table2.cell(13, 0).text = 'D21S11'
    table2.cell(13, 1).text = 'D21S11_1'
    table2.cell(13, 2).text = 'D21S11_2'
    table2.cell(13, 3).text = 'D21S11_3'
    table2.cell(13, 4).text = 'D21S11_4'
    table2.cell(13, 5).text = 'D21S11_bM'

    # table2.cell(14,

# Methodology:	Authenticity for all cell lines used in this study was performed using the Promega GenePrint 10 System and following the protocol described in ANSI/ATCC ASN-0002-2011.  The STR alleles were searched on the _dataset Database.
# Marker Characterization
# Marker	Allele 1	Allele 2	Allele 3	Allele 4	_dataset Best Match Profile
# D5S818	D5S818_1	D5S818_2	D5S818_3	D5S818_4	D5S818_bM
# D13S317	D13S317_1	D13S317_2	D13S317_3	D13S317_4	D13S317_bM
# D7S820	D7S820_1	D7S820_2	D7S820_3	D7S820_4	D7S820_bM
# D16S539	D16S539_1	D16S539_2	D16S539_3	D16S539_4	D16S539_bM
# vWA	vWA_1	vWA_2	vWA_3	vWA_4	vWA_bM
# TH01	TH01_1	TH01_2	TH01_3	TH01_4	TH01_bM
# AMEL	AMEL_1	AMEL_2	AMEL_3	AMEL_4	AMEL_bM
# TPOX	TPOX_1	TPOX_2	TPOX_3	TPOX_4	TPOX_bM
# CSF1PO	CSF1PO_1	CSF1PO_2	CSF1PO_3	CSF1PO_4	CSF1PO_bM
# D21S11	D21S11_1	D21S11_2	D21S11_3	D21S11_4	D21S11_bM





    doc.save('test.docx')


########################################################################################################################
# Helper Functions
########################################################################################################################

def set_cell_background(cell, fill):
    '''
    Sets the background color of a cell and centers the text
    Also bolds the text 
    '''
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_font_bold(cell)


def set_cell_font_bold(cell):
    '''
    Sets the font of a cell to bold
    '''
    for run in cell.paragraphs[0].runs:
        run.font.bold = True


if __name__ == "__main__":
    create_template()