# Author: Kenny Ma
# Contact: 626-246-2233 or kyoma17@gmail.com
# Date: 2022-12-23
# Version: 1.0
# Description: This script will take in an excel file from the GMapper program and perform the ClimaSTR Cell Line ID script on each sample
# Requirements: Selenium, Pandas, BeautifulSoup, docx, tkinter, Firefox, geckodriver.exe
# GekoDriver: Place geckodriver.exe in the same directory as this script
# Instructions: Run the script and select the input file from the GMapper program. The script will then perform the ClimaSTR Cell Line ID script on each sample and output a .docx file for each sample.
# The input file must have the following information: Sample Name, D5S818, D13S317, D7S820, D16S539, vWA, TH01, AMEL, TPOX, CSF1PO, D21S11

from tkinter import filedialog
from bs4 import BeautifulSoup
import docx
import tkinter
import pandas
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.options import Options
import time
import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
# Import the webdriver and time modules

def main():
    print("Welcome to the ClimaSTR Cell Line ID script" '\n' "Please select the input file from the GMapper Program")
    # get file path from user using tkinter file dialog
    tkinter.Tk().withdraw()
    # open file dialog in current directory and allow excel files only
    file_path = filedialog.askopenfilename(initialdir = ".", title = "Select file", filetypes = (("Select CellLine ID file","*.xlsx"),("all files","*.*")))

    # file_path = "C:/Users/kyo_m/Documents/Code/GP 10 input.xlsx"


    # read .xlsx into pandas dataframe
    df = pandas.read_excel(file_path)

    # Print dataframe and ask user to confirm correct file
    print(df)

    grouped_df = df.groupby("Sample Name")

    # for each group, perform the selenium script
    for each in grouped_df:
        sampleName = each[0]
        sampleDF = each[1]
        print("Processing " + sampleName + "...")
        ClimaSTRSearch(sampleName, sampleDF)

    input("Script complete. Press any key to exit")



def ClimaSTRSearch(sampleName, sampleDF):
    # Takes in a pandas dataframe of a single sample and performs the selenium script against Clima
    # Will return a pandas dataframe of the results of web scraping

    # Webpage input fields
    D5S818_list = ["D5S818", "D5S818_data1",
                   "D5S818_data2", "D5S818_data3", "D5S818_data4"]
    D13S317_list = ["D13S317", "D13S317_data1",
                    "D13S317_data2", "D13S317_data3", "D13S317_data4"]
    D7S820_list = ["D7S820", "D7S820_data1",
                   "D7S820_data2", "D7S820_data3", "D7S820_data4"]
    D16S539_list = ["D16S539", "D16S539_data1",
                    "D16S539_data2", "D16S539_data3", "D16S539_data4"]
    VWA_list = ["vWA", "VWA_data1", "VWA_data2", "VWA_data3", "VWA_data4"]
    TH01_list = ["TH01", "TH01_data1",
                 "TH01_data2", "TH01_data3", "TH01_data4"]
    Amelogenin_list = ["AMEL", "AMG_data1", "AMG_data2"]
    TPOX_list = ["TPOX", "TPOX_data1",
                 "TPOX_data2", "TPOX_data3", "TPOX_data4"]
    CSF1PO_list = ["CSF1PO", "CSF1PO_data1", "CSF1PO_data2",
                   "CSF1PO_data3", "CSF1PO_data4"]
    # D21S11_list = ["D21S11", "D21S11_data1", "D21S11_data2",
    #                "D21S11_data3", "D21S11_data4"]

    master_list = [D5S818_list, D13S317_list, D7S820_list, D16S539_list,
                   VWA_list, TH01_list, Amelogenin_list, TPOX_list, CSF1PO_list]

    # Create a new instance of the Firefox driver
    options = Options()
    options.binary_location = r'C:\Program Files\Mozilla Firefox\firefox.exe'
    driver = webdriver.Firefox(
        executable_path=r'C:\WebDrivers\geckodriver.exe', options=options)

    # Open a web browser and navigate to the website
    driver.get('http://bioinformatics.hsanmartino.it/clima2/index.php')

    # Wait for the page to load
    # time.sleep(1)

    # Input each allele into the webpage from the dataframe
    for each in master_list:
        marker = each[0]
        row = sampleDF.loc[sampleDF["Markers"] == marker]

        for i in range(1, len(each)):
            # skip nan values
            if pandas.isna(row["Allele" + str(i)].values[0]):
                continue

            allele = driver.find_element("id", each[i])

            allele.send_keys(row["Allele" + str(i)].values[0])

    # enter email and country and submit
    email = driver.find_element("id", "usr_email")
    email.send_keys("kyoma17@gmail.com")
    country = driver.find_element("id", "usr_country")
    country.send_keys("United States")

    # Find the submit button
    submit_button = driver.find_element(
        By.XPATH, "//input[@type='submit' and @value='submit']")

    # Click the submit button
    submit_button.click()

    # Wait for the page to load
    time.sleep(1)

    # retreive table from webpage and convert to pandas dataframe
    table = driver.find_element(By.XPATH, "(//table)[3]")
    html_table = table.get_attribute("outerHTML")
    soup = BeautifulSoup(html_table, "html.parser")

    # find the table
    souptable = soup.find("table")
    
    rows = souptable.find_all("tr")
    data = []
    for row in rows:
        cells = row.find_all("td")
        values = [cell.text for cell in cells]
        data.append(values)
    tableDF = pandas.DataFrame(data)

    PandasTableDF = pandas.read_html(html_table)[0]
    PandasTableDF.reset_index(drop=True, inplace=True)
    tableDF.columns = PandasTableDF.columns.get_level_values(0)

    # select the highest match
    bestMatched = tableDF.iloc[2]

    # Close the browser
    driver.quit()

    # Transfer the data to the docx template
    document = docx.Document('CellLineTemplateGP10.docx')
    replacementsDictionary = {
        # Data from Input
        "_SAMPLE_NAME": sampleName,
        "D5S818_1": sampleDF.loc[sampleDF["Markers"] == "D5S818"]["Allele1"].values[0],
        "D5S818_2": sampleDF.loc[sampleDF["Markers"] == "D5S818"]["Allele2"].values[0],
        "D5S818_3": sampleDF.loc[sampleDF["Markers"] == "D5S818"]["Allele3"].values[0],
        "D5S818_4": sampleDF.loc[sampleDF["Markers"] == "D5S818"]["Allele4"].values[0],
        "D13S317_1": sampleDF.loc[sampleDF["Markers"] == "D13S317"]["Allele1"].values[0],
        "D13S317_2": sampleDF.loc[sampleDF["Markers"] == "D13S317"]["Allele2"].values[0],
        "D13S317_3": sampleDF.loc[sampleDF["Markers"] == "D13S317"]["Allele3"].values[0],
        "D13S317_4": sampleDF.loc[sampleDF["Markers"] == "D13S317"]["Allele4"].values[0],
        "D7S820_1": sampleDF.loc[sampleDF["Markers"] == "D7S820"]["Allele1"].values[0],
        "D7S820_2": sampleDF.loc[sampleDF["Markers"] == "D7S820"]["Allele2"].values[0],
        "D7S820_3": sampleDF.loc[sampleDF["Markers"] == "D7S820"]["Allele3"].values[0],
        "D7S820_4": sampleDF.loc[sampleDF["Markers"] == "D7S820"]["Allele4"].values[0],
        "D16S539_1": sampleDF.loc[sampleDF["Markers"] == "D16S539"]["Allele1"].values[0],
        "D16S539_2": sampleDF.loc[sampleDF["Markers"] == "D16S539"]["Allele2"].values[0],
        "D16S539_3": sampleDF.loc[sampleDF["Markers"] == "D16S539"]["Allele3"].values[0],
        "D16S539_4": sampleDF.loc[sampleDF["Markers"] == "D16S539"]["Allele4"].values[0],
        "vWA_1": sampleDF.loc[sampleDF["Markers"] == "vWA"]["Allele1"].values[0],
        "vWA_2": sampleDF.loc[sampleDF["Markers"] == "vWA"]["Allele2"].values[0],
        "vWA_3": sampleDF.loc[sampleDF["Markers"] == "vWA"]["Allele3"].values[0],
        "vWA_4": sampleDF.loc[sampleDF["Markers"] == "vWA"]["Allele4"].values[0],
        "TH01_1": sampleDF.loc[sampleDF["Markers"] == "TH01"]["Allele1"].values[0],
        "TH01_2": sampleDF.loc[sampleDF["Markers"] == "TH01"]["Allele2"].values[0],
        "TH01_3": sampleDF.loc[sampleDF["Markers"] == "TH01"]["Allele3"].values[0],
        "TH01_4": sampleDF.loc[sampleDF["Markers"] == "TH01"]["Allele4"].values[0],
        "AMEL_1": sampleDF.loc[sampleDF["Markers"] == "AMEL"]["Allele1"].values[0],
        "AMEL_2": sampleDF.loc[sampleDF["Markers"] == "AMEL"]["Allele2"].values[0],
        "AMEL_3": sampleDF.loc[sampleDF["Markers"] == "AMEL"]["Allele3"].values[0],
        "AMEL_4": sampleDF.loc[sampleDF["Markers"] == "AMEL"]["Allele4"].values[0],
        "TPOX_1": sampleDF.loc[sampleDF["Markers"] == "TPOX"]["Allele1"].values[0],
        "TPOX_2": sampleDF.loc[sampleDF["Markers"] == "TPOX"]["Allele2"].values[0],
        "TPOX_3": sampleDF.loc[sampleDF["Markers"] == "TPOX"]["Allele3"].values[0],
        "TPOX_4": sampleDF.loc[sampleDF["Markers"] == "TPOX"]["Allele4"].values[0],
        "CSF1PO_1": sampleDF.loc[sampleDF["Markers"] == "CSF1PO"]["Allele1"].values[0],
        "CSF1PO_2": sampleDF.loc[sampleDF["Markers"] == "CSF1PO"]["Allele2"].values[0],
        "CSF1PO_3": sampleDF.loc[sampleDF["Markers"] == "CSF1PO"]["Allele3"].values[0],
        "CSF1PO_4": sampleDF.loc[sampleDF["Markers"] == "CSF1PO"]["Allele4"].values[0],
        "D21S11_1": sampleDF.loc[sampleDF["Markers"] == "D21S11"]["Allele1"].values[0],
        "D21S11_2": sampleDF.loc[sampleDF["Markers"] == "D21S11"]["Allele2"].values[0],
        "D21S11_3": sampleDF.loc[sampleDF["Markers"] == "D21S11"]["Allele3"].values[0],
        "D21S11_4": sampleDF.loc[sampleDF["Markers"] == "D21S11"]["Allele4"].values[0],

        # Data from the Results highest scoring match
        "_bMatchScore": bestMatched["% Match"],
        "_bMatchName": bestMatched["Name"],
        "_bMatchCellLineNo": bestMatched["Cat. No."],

        "D5S818_bM": bestMatched["D5S818"],
        "D13S317_bM": bestMatched["D13S317"],
        "D7S820_bM": bestMatched["D7S820"],
        "D16S539_bM": bestMatched["D16S539"],
        "vWA_bM": bestMatched["VWA"],
        "TH01_bM": bestMatched["TH01"],
        "AMEL_bM": bestMatched["AMG"],
        "TPOX_bM": bestMatched["TPOX"],
        "CSF1PO_bM": bestMatched["CSF1PO"],
        "D21S11_bM": bestMatched["D21S11"],
    }

    # Iterate through the tables in the document and replace the text from the dictionary
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in replacementsDictionary.keys():
                        if key in paragraph.text:
                            # Enter the replacement text in the document, if the value is None or nan, then replace with an empty string
                            if replacementsDictionary[key] is None or pandas.isna(replacementsDictionary[key]):
                                paragraph.text = paragraph.text.replace(
                                    key, "")
                            else:
                                paragraph.text=paragraph.text.replace(
                                    key, str(replacementsDictionary[key]))


    # Save the modified document
    document.save(sampleName + ".docx")
    print("Done with " + sampleName)
    
main()
