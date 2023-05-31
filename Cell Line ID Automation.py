# Author: Kenny Ma
# Contact: 626-246-2233 or kyoma17@gmail.com
# Date: 2023-May-29
version = 3.5
# Description: This script will take in an excel file from the GMapper program and perform the ClimaSTR Cell Line ID script on each sample
# The script will then consolidate the results into a single .docx file
# Requirements: Selenium, Pandas, BeautifulSoup, docx, tkinter, Firefox, geckodriver.exe
# there is a Requirements.txt file that can be used to install the required modules
# Use pip install -r Requirements.txt to install the required modules
# GekoDriver: Place geckodriver.exe in the same directory as this script
# Instructions: Run the script and select the input file from the GMapper program. The script will then perform the ClimaSTR Cell Line ID script on each sample and output a .docx file for each sample.
# The input file must have the following information: Sample Name, D5S818, D13S317, D7S820, D16S539, vWA, TH01, AMEL, TPOX, CSF1PO, D21S11

# You must have Word installed on your computer with the Trust Center settings set to "Enable all macros"
from tkinter import filedialog
import tkinter as tk
from tkinter import ttk
from bs4 import BeautifulSoup
import threading
import docx
import tkinter
from tkinter import messagebox
import pandas
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver import ActionChains
from selenium.webdriver.firefox.options import Options
import time
import warnings
import win32com.client
import os
import psutil
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)


def main():
    close_word_processes()
    print("Welcome to the ClimaSTR Cell Line ID script" '\n' "Please select the input file from the GMapper Program")
    print("version: " + str(version))
    # get file path from user using tkinter file dialog
    tkinter.Tk().withdraw()
    # open file dialog in current directory and allow excel files only
    file_path = filedialog.askopenfilename(initialdir=".", title="Select file", filetypes=(
        ("Select CellLine ID file", "*.xlsx"), ("all files", "*.*")))

    # Load the client list excel file
    selected_client, reference_number = SelectClient()

    client_database = pd.read_excel("CellLineClients.xlsx")

    # Get the client row from the client database
    selected_client_info = client_database.loc[client_database["Nickname"]
                                               == selected_client]

    # file_path = "C:/Users/kyo_m/Documents/Code/GP 10 input.xlsx"

    # If the output folder does not exist, create it
    if not os.path.exists("CellLineTEMP"):
        os.makedirs("CellLineTEMP")

    # read .xlsx into pandas dataframe
    df = pandas.read_excel(file_path)

    sampleList = df["Sample Name"].unique()

    # Print name of file to console
    print("Loaded file: " + file_path + '\n' "Processing..." + '\n')

    grouped_df = df.groupby("Sample Name")

    sample_counter = 0
    sample_order = []

    # for each group, perform the selenium script
    for each in grouped_df:
        sample_counter += 1

        # sample_counter = "#"
        # if Test Name is geneprint24

        sampleName = each[0]
        sample_order.append(sampleName)
        sampleDF = each[1]
        testName = each[1]["Test Name"].values

        if "GenePrint_24_POP7_Panels_v1.0" in testName:
            print("Processing GP24 " + sampleName + "...")
        else:
            print("Processing GP10 " + sampleName + "...")

        expasy_results = ExpasySTRSearch(sampleName, sampleDF, sample_counter)
        clima_results = ClimaSTRSearch(sampleName, sampleDF, sample_counter)

        # Combine the results from ClimaSTR and ExpasySTR

        results = clima_results + expasy_results

        # Display the best matched samples to the user and ask for user input
        selectSample(results)
        close_microsoft_word()

    # run vba macro to save all word documents

    finalReport(sample_order, selected_client_info, reference_number)

    # input("Cell Line ID script has finished running" '\n' "Press ENTER to exit")
    show_done_window()

def close_word_processes():
    for process in psutil.process_iter(['pid', 'name']):
        if process.info['name'] == 'WINWORD.EXE':  # Check for the Word process name
            try:
                process.kill()  # Terminate the Word process
                print(f"Closed Word process with PID: {process.info['pid']}")
            except psutil.AccessDenied:
                print(f"Access denied to terminate Word process with PID: {process.info['pid']}")


def finalReport(listOfSamples, clientInfo, reference_number):
    # Create Replace Dictionary for the final report
    # This dictionary is for the Header in the Paperwork
    replaceDict = {"_PIName": clientInfo["PIName"].values[0],
                   "_ClientName": clientInfo["ClientName"].values[0],
                   "_ClientEmail": clientInfo["ClientEmail"].values[0],
                   "_ClientPhoneNumber": clientInfo["ClientPhoneNumber"].values[0],
                   "_Institution": clientInfo["Institution"].values[0],
                   "_ReferenceNumber": reference_number,
                   "_Date": time.strftime("%m/%d/%Y"),
                   "_SampleNumber": len(listOfSamples),
                   "_Batches": "1",
                   }


    # Open the Header and Footer template
    combined_document = docx.Document('ClientTemplate.docx')

    # iterate through the other documents and add them to the combined document

    # add the content of each document to the combined document keep the formatting each document
    # add a page break to the end of each document except the last one
    for doc in listOfSamples:
        temp_doc = docx.Document("CellLineTEMP/" + doc + ".docx")
        for element in temp_doc.element.body:
            combined_document.element.body.append(element)
        # combined_document.add_page_break()

    # Show header and footer from first page in all pages
    for section in combined_document.sections:
        section.different_first_page_header_footer = False

    # Replace spaces in client name with underscores
    clientFileName = clientInfo["Nickname"].values[0].replace(" ", "_")

    # Save the final report in CellLineOutput folder
    report_name = "CellLineOutput/" + clientFileName + "_Cell_Line_ID_" + reference_number + ".docx" 
    combined_document.save(report_name)

    # Convert the Python dictionary to a VBA dictionary
    vba_dict = win32com.client.Dispatch("Scripting.Dictionary")

    # Iterate through the Python dictionary and add the key/value pairs to the VBA dictionary
    for key in replaceDict:
        vba_dict.Add(key, replaceDict[key])

    # get the absolute path of the word document
    report_name = os.path.abspath(report_name)

    # open word document
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = 0

    doc = word.Documents.Open(report_name, ReadOnly=1  )

    # Load the VBA code from the file
    with open("CellLineOutputVBA.bas", "r") as f:
        vbaCode = f.read()


    # Inject vba script into word document
    doc.VBProject.VBComponents.Add(1).CodeModule.AddFromString(vbaCode)

    # run macro
    word.Run("ReplaceHeaderKeyword", vba_dict)

    # save and close word document
    doc.Save()
    doc.Close()

def close_microsoft_word():
    # close word application if it is open
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Quit()
    except:
        pass


def selectSample(bestMatchedSamples):
    # Display the best matched samples to the user and ask for user input
    window = tk.Tk()
    window.title("Select Result")
    window.columnconfigure(0, minsize=250, weight=1)
    window.rowconfigure([0, 1], minsize=100, weight=1)

    tree = ttk.Treeview(window)

    columns = ("name", "_dataset", "_bMatchScore", "_bMatchName", "_bMatchCellLineNo", "D5S818_bM", "D13S317_bM",
               "D7S820_bM", "D16S539_bM", "vWA_bM", "TH01_bM", "AMEL_bM", "TPOX_bM", "CSF1PO_bM", "D21S11_bM")

    # Create the treeview
    tree = ttk.Treeview(window, columns=columns, show='headings')

    # Set the column widths
    for col in columns:
        tree.column(col, width=100, anchor=tk.W)

    # Set the column headings
    for col, heading in enumerate(columns):
        tree.heading(col, text=heading, anchor=tk.W)

    # unpack Query information from sampleDF and add to treeview

    # unpack list of dictionaries into a dictionary of dictionaries

    # This section is only for show  to the user
    treeviewDictionary = {}

    climaCounter = 1
    expasyCounter = 1

    for i, data in enumerate(bestMatchedSamples):
        if data["website"] == "Clima2":
            treeviewDictionary[f"Clima{climaCounter}"] = data
            climaCounter += 1
        else:
            treeviewDictionary[f"Expasy{expasyCounter}"] = data
            expasyCounter += 1

    for i, (name, data) in enumerate(treeviewDictionary.items()):
        tree.insert("", i, text=name, values=(name, data["_dataset"], data["_bMatchScore"], data["_bMatchName"], data["_bMatchCellLineNo"], data["D5S818_bM"],
                    data["D13S317_bM"], data["D7S820_bM"], data["D16S539_bM"], data["vWA_bM"], data["TH01_bM"], data["AMEL_bM"], data["TPOX_bM"], data["CSF1PO_bM"], data["D21S11_bM"]))


# Create a submit button, close window when clicked


    def submit():
        # Get the selected item
        item = tree.selection()[0]
        # Print the Name of the selected item
        selection = tree.item(item, "values")[0]
        print("Selected result: ", selection, " with ",
              treeviewDictionary[selection]["_bMatchScore"], " match score")
        # Pull selected item from dictionary and write to template
        fillTemplate(treeviewDictionary[selection])

        # Close the window
        window.destroy()
        window.quit()

    submit_button = tk.Button(window, text='SELECT RESULT', command=submit)
    tree.pack()
    submit_button.pack()

    # Run the main loop
    window.mainloop()


############################################################################################################
def ExpasySTRSearch(sampleName, sampleDF, sampleNumber):
    # Takes in a pandas dataframe of a single sample and performs the selenium script against Expasy
    # Will return a pandas dataframe of the results of web scraping

    # Webpage input fields
    Amelogen = ["AMEL", "input-Amelogenin"]
    CSF1PO = ["CSF1PO", "input-CSF1PO"]
    D2S133 = ["D2S133", "input-D2S133"]
    D3S135 = ["D3S1358", "input-D3S1358"]
    D5S818 = ["D5S818", "input-D5S818"]
    D7S820 = ["D7S820", "input-D7S820"]
    D8S1179 = ["D8S1179", "input-D8S1179"]
    D13S317 = ["D13S317", "input-D13S317"]
    D16S539 = ["D16S539", "input-D16S539"]
    D18S51 = ["D18S51", "input-D18S51"]
    D19S433 = ["D19S433", "input-D19S433"]
    D21S11 = ["D21S11", "input-D21S11"]
    FGA = ["FGA", "input-FGA"]
    PentaD = ["Penta D", "input-Penta_D"]
    PentaE = ["Penta E", "input-Penta_E"]
    TH01 = ["TH01", "input-TH01"]
    TPOX = ["TPOX", "input-TPOX"]
    vWA = ["vWA", "input-vWA"]
    D2S1338 = ["D2S1338", "input-D2S1338"]

    D10S1248 = ["D10S1248", "input-D10S1248"]
    D12S391 = ["D12S391", "input-D12S391"]
    D22S1045 = ["D22S1045", "input-D22S1045"]
    D2S441 = ["D2S441", "input-D2S441"]
    D1S1656 = ["D1S1656", "input-D1S1656"]
    DYS391 = ["DYS391", "input-DYS391"]

    master_list = [Amelogen, CSF1PO, D2S133, D3S135, D5S818, D7S820, D8S1179,
                   D13S317, D16S539, D18S51, D19S433, D21S11, FGA, PentaD, PentaE,
                   TH01, TPOX, vWA, D10S1248, D12S391, D22S1045, D2S441, D1S1656, DYS391, D2S1338]

    # Create a new instance of the Firefox driver
    options = Options()
    options.binary_location = r'C:\Program Files\Mozilla Firefox\firefox.exe'
    driver = webdriver.Firefox(
        executable_path=r'C:\WebDrivers\geckodriver.exe', options=options)

    # go to the Expasy STR website
    driver.get("https://www.cellosaurus.org/str-search/")
    time.sleep(1)

    # Click Checkboxes if GP24
    if "GenePrint_24_POP7_Panels_v1.0" in sampleDF["Test Name"].values:
        clickList = ["check-D10S1248", "check-D12S391", "check-D22S1045",
                     "check-D2S441", "check-D1S1656", "check-DYS391"]
        for each in clickList:
            ActionChains(driver).move_to_element(
                driver.find_element("id", each)).click().perform()

    # Input each allele into the webpage from the dataframe
    for each in master_list:
        marker = each[0]
        row = sampleDF.loc[sampleDF["Markers"] == marker]

        # combine Allele 1-4 from row into a single string, skip if NaN
        alleles = []
        for i in range(1, 5):
            allele = row[f"Allele{i}"].values
            if len(allele) > 0 and not pandas.isnull(allele[0]):
                alleles.append(str(allele[0]))

        allele = ",".join(alleles) if len(alleles) > 0 else ""

        # Input the allele into the webpage
        if len(allele) > 0:
            driver.find_element("id", each[1]).send_keys(allele)

    # Click the search button
    driver.find_element("id", "search").click()

    # Wait for the results to load
    time.sleep(2)

    # if no results, return empty list of empty dictionary

    warning_element = driver.find_element("id", "warning")
    display_value = warning_element.value_of_css_property('display')
    if display_value == 'block':
        print("No results for " + sampleName + " from Expasy STR Search")
        driver.quit()
        return []

    # if "Warning:" in driver.page_source or "The query returned no results." in driver.page_source:

    table = driver.find_element("id", "table-results")
    html = table.get_attribute("outerHTML")
    soup = BeautifulSoup(html, "html.parser")
    rows = soup.find_all("tr")
    header_row = rows[0]
    column_names = [th.text for th in header_row.find_all("th")]

    # Store the data in a list of lists
    data = []
    for row in rows:
        cells = row.find_all("td")
        data.append([cell.text for cell in cells])

    tableDF = pd.DataFrame(data, columns=column_names)
    # input("Press Enter to continue...")
    # Close the browser
    driver.quit()

    # number of results to return
    number_of_results = 10

    def get_best_match(tableDF, i):
        try :
            bestMatched.append(tableDF.iloc[i])
        
        except IndexError:
            print(tableDF)
            quit()

    bestMatched = []
    threads = []
    for i in range(2, number_of_results + 2):
        thread = threading.Thread(target=get_best_match, args=(tableDF, i))
        threads.append(thread)
        thread.start()

    # Wait for all threads to complete
    for thread in threads:
        thread.join()

    replacementsDictionaries = []
    # Create a list to store the threads
    threads = []

    # Create a function to run in a separate thread
    def generate_replacement_dictionary_thread(sampleName, sampleDF, bestMatched, prefix):
        replacementsDictionaries.append(generateReplacementDictionary(
            sampleName, sampleDF, bestMatched, prefix, sampleNumber))

    # Create a thread for each call to generateReplacementDictionary
    for i in range(len(bestMatched)):
        t = threading.Thread(target=generate_replacement_dictionary_thread, args=(
            sampleName, sampleDF, bestMatched[i], "Expasy"))
        threads.append(t)

    # Start the threads
    for t in threads:
        t.start()

    # Wait for the threads to finish
    for t in threads:
        t.join()

    return replacementsDictionaries


def ClimaSTRSearch(sampleName, sampleDF, sampleNumber):
    # Takes in a pandas dataframe of a single sample and performs the selenium script against Clima
    # Will return a pandas dataframe of the results of web scraping

    # Webpage input fields
    D5S818_list = ["D5S818", "D5S818_data1",
                   "D5S818_data2", "D5S818_data3", "D5S818_data4"]
    D13S317_list = ["D13S317", "D13S317_data1",
                    "D13S317_data2", "D13S317_data3", "D13S317_data4"]
    D7S820_list = ["D7S820", "D7S820_data1",
                   "D7S820_data2", "D7S820_data3", "D7S820_data4"]
    D16S539_list = ["D16S539", "D16S539_data1",
                    "D16S539_data2", "D16S539_data3", "D16S539_data4"]
    VWA_list = ["vWA", "VWA_data1", "VWA_data2", "VWA_data3", "VWA_data4"]
    TH01_list = ["TH01", "TH01_data1",
                 "TH01_data2", "TH01_data3", "TH01_data4"]
    Amelogenin_list = ["AMEL", "AMG_data1", "AMG_data2"]
    TPOX_list = ["TPOX", "TPOX_data1",
                 "TPOX_data2", "TPOX_data3", "TPOX_data4"]
    CSF1PO_list = ["CSF1PO", "CSF1PO_data1", "CSF1PO_data2",
                   "CSF1PO_data3", "CSF1PO_data4"]
    # D21S11_list = ["D21S11", "D21S11_data1", "D21S11_data2",
    #                "D21S11_data3", "D21S11_data4"]

    master_list = [D5S818_list, D13S317_list, D7S820_list, D16S539_list,
                   VWA_list, TH01_list, Amelogenin_list, TPOX_list, CSF1PO_list]

    # Create a new instance of the Firefox driver
    options = Options()
    options.binary_location = r'C:\Program Files\Mozilla Firefox\firefox.exe'
    driver = webdriver.Firefox(
        executable_path=r'C:\WebDrivers\geckodriver.exe', options=options)

    # Open a web browser and navigate to the website
    driver.get('http://bioinformatics.hsanmartino.it/clima2/index.php')

    # Wait for the page to load
    # time.sleep(1)

    # Input each allele into the webpage from the dataframe
    for each in master_list:
        marker = each[0]
        row = sampleDF.loc[sampleDF["Markers"] == marker]

        for i in range(1, len(each)):
            # skip nan values
            if pandas.isna(row["Allele" + str(i)].values[0]):
                continue

            allele = driver.find_element("id", each[i])

            allele.send_keys(row["Allele" + str(i)].values[0])

    # enter email and country and submit
    email = driver.find_element("id", "usr_email")
    email.send_keys("info@laragen.com")
    country = driver.find_element("id", "usr_country")
    country.send_keys("United States")

    # Find the submit button
    submit_button = driver.find_element(
        By.XPATH, "//input[@type='submit' and @value='submit']")

    # Click the submit button
    submit_button.click()

    # Wait for the page to load

    # retreive table from webpage and convert to pandas dataframe
    table = driver.find_element(By.XPATH, "(//table)[3]")
    html_table = table.get_attribute("outerHTML")
    soup = BeautifulSoup(html_table, "html.parser")

    # find the table
    souptable = soup.find("table")
    rows = souptable.find_all("tr")

    data = []
    for row in rows:
        cells = row.find_all("td")
        values = [cell.text for cell in cells]
        data.append(values)
    tableDF = pandas.DataFrame(data)

    PandasTableDF = pandas.read_html(html_table)[0]
    PandasTableDF.reset_index(drop=True, inplace=True)
    tableDF.columns = PandasTableDF.columns.get_level_values(0)

    # input("Press Enter to continue...")
    # # Close the browser
    driver.quit()

    # number of results to return
    # check how many rows in the table
    numberOfHits = len(tableDF.index) - 2

    if numberOfHits < 10:
        number_of_results = numberOfHits
    else:
        number_of_results = 10

    def get_best_match(tableDF, i):
        # print how many rows in the table
        bestMatched.append(tableDF.iloc[i])

    bestMatched = []
    threads = []
    for i in range(2, number_of_results + 2):
        thread = threading.Thread(target=get_best_match, args=(tableDF, i))
        threads.append(thread)
        thread.start()

    # Wait for all threads to complete
    for thread in threads:
        thread.join()

    replacementsDictionaries = []

    # Create a list to store the threads
    threads = []

    # Create a function to run in a separate thread
    def generate_replacement_dictionary_thread(sampleName, sampleDF, bestMatched, prefix):
        replacementsDictionaries.append(generateReplacementDictionary(
            sampleName, sampleDF, bestMatched, prefix, sampleNumber))

    # Create a thread for each call to generateReplacementDictionary
    for i in range(number_of_results):
        t = threading.Thread(target=generate_replacement_dictionary_thread, args=(
            sampleName, sampleDF, bestMatched[i], "Clima"))
        threads.append(t)

    # Start the threads
    for t in threads:
        t.start()

    # Wait for the threads to finish
    for t in threads:
        t.join()

    return replacementsDictionaries

########################################################################################################################


def fillTemplate(replacementsDictionary):
    # Helper function to fill the template with the data from the dictionary
    sampleName = replacementsDictionary["_SAMPLE_NAME"]

    # Determine which template to use based on the test name

    if replacementsDictionary["test"].values[0] == "GenePrint_24_POP7_Panels_v1.0":
        document = docx.Document('CellLineTemplateGP24.docx')
    elif replacementsDictionary["test"].values[0] == "GenePrint_10_v1.1":
        document = docx.Document('CellLineTemplateGP10.docx')

    # Iterate through the tables in the document and replace the text from the dictionary, Preserve the font attributes
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:

                    for run in paragraph.runs:

                        for key in replacementsDictionary.keys():
                            if key in run.text:
                                # If the key is any of the special keys, peform a special action to color the text

                                # Flag to color the text red
                                keys_to_check = ["D5S818_bM", "D13S317_bM", "D7S820_bM", "D16S539_bM",
                                                 "vWA_bM", "TH01_bM", "AMEL_bM", "TPOX_bM", "CSF1PO_bM", "D21S11_bM"]

                                if key in keys_to_check:

                                    # run color replacement
                                    NumbersToColorRed = []
                                    # Get the value of the key and seperate by comma
                                    dictionaryValue = replacementsDictionary[key].split(
                                        ",")

                                    textToReplaceWith = replacementsDictionary[key]

                                    alleleLookupValue1 = replacementsDictionary[key[:-3] + "_1"]
                                    alleleLookupValue2 = replacementsDictionary[key[:-3] + "_2"]
                                    alleleLookupValue3 = replacementsDictionary[key[:-3] + "_3"]
                                    alleleLookupValue4 = replacementsDictionary[key[:-3] + "_4"]

                                    ValuesToCheck = []

                                    for allele in [alleleLookupValue1, alleleLookupValue2, alleleLookupValue3, alleleLookupValue4]:
                                        # if nan skip
                                        if not pandas.isna(allele):
                                            ValuesToCheck.append(allele)

                                    for value in dictionaryValue:
                                        if value in ValuesToCheck:
                                            NumbersToColorRed.append(value)

                                # run Normal replacement
                                # Save the current font attributes
                                font_name = run.font.name
                                font_size = run.font.size
                                font_bold = run.font.bold
                                font_italic = run.font.italic

                                # Replace the text
                                if replacementsDictionary[key] is None or pandas.isna(replacementsDictionary[key]):
                                    run.text = run.text.replace(key, "")
                                else:
                                    run.text = run.text.replace(
                                        key, str(replacementsDictionary[key]))

                                # Apply the saved font attributes to the Font object
                                run.font.name = font_name
                                run.font.size = font_size
                                run.font.bold = font_bold
                                run.font.italic = font_italic

    # Second pass: Iterate through the tables in the document and replace the text from the dictionary
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in replacementsDictionary.keys():
                        if key in paragraph.text:
                            # Enter the replacement text in the document, if the value is None or nan, then replace with an empty string
                            if replacementsDictionary[key] is None or pandas.isna(replacementsDictionary[key]):
                                paragraph.text = paragraph.text.replace(
                                    key, "")
                            else:
                                paragraph.text = paragraph.text.replace(
                                    key, str(replacementsDictionary[key]))

    # Inject and run VBA code to the document

    # Save the modified document to the CellLineTEMP folder in the current directory
    document.save('CellLineTEMP/' + sampleName + '.docx')

    injectAndRunRedCodeVBA('CellLineTEMP/' + sampleName + '.docx')

    print("Done with " + sampleName)
    # Print line empty line
    print("")


def injectAndRunRedCodeVBA(fileName):
    # Inject and run VBA code to the document

    # Current Directory
    currentDirectory = os.getcwd()

    filePath = os.path.join(currentDirectory, fileName)

    # Open the document with win32com turn on the visible mode
    word = win32com.client.DispatchEx("Word.Application")

    # Open the document in read only mode
    doc = word.Documents.Open(filePath, ReadOnly=True)

    # Load the VBA code from the file
    with open("CellLineRed.bas", "r") as f:
        vbaCode = f.read()

    # Inject the VBA code to the document
    doc.VBProject.VBComponents.Add(1).CodeModule.AddFromString(vbaCode)

    # Run the VBA code
    word.Run("ChangeMatchingToRed")

    # Save the document 
    doc.Save()

    # Close the document
    doc.Close()

    # Quit the word application
    word.Quit()

########################################################################################################################


def generateReplacementDictionary(sampleName, sampleDF, bestMatched, website, sampleNumber):
    # Generate Replacement Dictionary for the Template

    # Clima dictionary for the template
    if website == "Clima":
        replacementsDictionary = {
            # Main Info
            "_SAMPLE_NAME": sampleName,
            "_sampleNumber": sampleNumber,
            "website": "Clima2",
            "test": sampleDF["Test Name"],

            # Data from the Results highest scoring match
            "_dataset": bestMatched["Dataset"],
            "_bMatchScore": bestMatched["% Match"],
            "_bMatchName": bestMatched["Name"],
            "_bMatchCellLineNo": bestMatched["Cat. No."],

            "D5S818_bM": bestMatched["D5S818"],  # Marker 1
            "D13S317_bM": bestMatched["D13S317"],  # Marker 2
            "D7S820_bM": bestMatched["D7S820"],  # Marker 3
            "D16S539_bM": bestMatched["D16S539"],  # Marker 4
            "vWA_bM": bestMatched["VWA"],  # Marker 5
            "TH01_bM": bestMatched["TH01"],  # Marker 6
            "AMEL_bM": bestMatched["AMG"],  # Marker 7
            "TPOX_bM": bestMatched["TPOX"],  # Marker 8
            "CSF1PO_bM": bestMatched["CSF1PO"],  # Marker 9
            "D21S11_bM": bestMatched["D21S11"],  # Marker 10
        }
        if sampleDF["Test Name"].values[0] == "GenePrint_24_POP7_Panels_v1.0":
            replacementDictionaryGP24 = {
                "D10S1248_bM": "",  # Marker 11
                "D12S391_bM": "",  # Marker 12
                "D18S51_bM": bestMatched["D18S51"],  # Marker 13
                "D19S433_bM": bestMatched["D19S433"],  # Marker 14
                "D1S1656_bM": "",  # Marker 15
                "D22S1045_bM": "",  # Marker 16
                "D2S1338_bM": bestMatched["D2S1338"],  # Marker 17
                "D2S441_bM": "",  # Marker 18
                "D3S1358_bM": bestMatched["D3S1358"],  # Marker 19
                "D8S1179_bM": bestMatched["D8S1179"],  # Marker 20
                "DYS391_bM": "",  # Marker 21
                "FGA_bM": bestMatched["FGA"],  # Marker 22
                "Penta E_bM": bestMatched["PentaD"],  # Marker 23
                "Penta D_bM": bestMatched["PentaE"],  # Marker 24
            }
            replacementsDictionary.update(replacementDictionaryGP24)

    # Expasy dictionary for the template
    elif website == "Expasy":
        replacementsDictionary = {

            "_SAMPLE_NAME": sampleName,
            "website": "Expasy",
            "test": sampleDF["Test Name"],

            # Data from the Results highest scoring match
            "_dataset": "Expasy",
            "_bMatchScore": bestMatched["Score"],
            "_bMatchName": bestMatched["Accession"],
            "_bMatchCellLineNo": bestMatched["Name"],

            "D5S818_bM": bestMatched["D5S818"],  # Marker 1
            "D13S317_bM": bestMatched["D13S317"],  # Marker 2
            "D7S820_bM": bestMatched["D7S820"],  # Marker 3
            "D16S539_bM": bestMatched["D16S539"],  # Marker 4
            "vWA_bM": bestMatched["vWA"],  # Marker 5
            "TH01_bM": bestMatched["TH01"],  # Marker 6
            "AMEL_bM": bestMatched["Amel"],  # Marker 7
            "TPOX_bM": bestMatched["TPOX"],  # Marker 8
            "CSF1PO_bM": bestMatched["CSF1PO"],  # Marker 9
            "D21S11_bM": bestMatched["D21S11"],  # Marker 10
        }

        # If GenePrint24 Add the extra markers
        if sampleDF["Test Name"].values[0] == "GenePrint_24_POP7_Panels_v1.0":
            replacementDictionaryGP24 = {
                "D10S1248_bM": bestMatched["D10S1248"],  # Marker 11
                "D12S391_bM": bestMatched["D12S391"],  # Marker 12
                "D18S51_bM": bestMatched["D18S51"],  # Marker 13
                "D19S433_bM": bestMatched["D19S433"],  # Marker 14
                "D1S1656_bM": bestMatched["D1S1656"],  # Marker 15
                "D22S1045_bM": bestMatched["D22S1045"],  # Marker 16
                "D2S1338_bM": bestMatched["D2S1338"],  # Marker 17
                "D2S441_bM": bestMatched["D2S441"],  # Marker 18
                "D3S1358_bM": bestMatched["D3S1358"],  # Marker 19
                "D8S1179_bM": bestMatched["D8S1179"],  # Marker 20
                "DYS391_bM": bestMatched["DYS391"],  # Marker 21
                "FGA_bM": bestMatched["FGA"],  # Marker 22
                "Penta E_bM": bestMatched["Penta E"],  # Marker 23
                "Penta D_bM": bestMatched["Penta D"],  # Marker 24
            }
            replacementsDictionary.update(replacementDictionaryGP24)

    # Replacement Dictionary for the Template from  Data Input
    markers = ["D5S818", "D13S317", "D7S820", "D16S539",
               "vWA", "TH01", "AMEL", "TPOX", "CSF1PO", "D21S11"]

    # If GenePrint24 Add the extra markers
    if sampleDF["Test Name"].values[0] == "GenePrint_24_POP7_Panels_v1.0":
        markers.extend(["D10S1248", "D12S391", "D18S51", "D19S433", "D1S1656", "D22S1045",
                        "D2S1338", "D2S441", "D3S1358", "D8S1179", "DYS391", "FGA", "Penta E", "Penta D"])

    allele_columns = ["Allele1", "Allele2", "Allele3", "Allele4"]

    template_dict = {}

    for marker in markers:
        for i, allele_column in enumerate(allele_columns, start=1):
            key = f"{marker}_{i}"

            value = sampleDF.loc[sampleDF["Markers"]
                                 == marker][allele_column].values[0]
            template_dict[key] = value

    replacementsDictionary.update(template_dict)

    return replacementsDictionary


def SelectClient():
    # Selects the client from the listbox and returns the order number
    print("Select Client from the listbox and enter the order number")

    # Load Client Data from Excel File and create a dataframe
    client_database = pd.read_excel("CellLineClients.xlsx")
    client_list = client_database["Nickname"].tolist()

    selected_item = ""
    order_number = ""

    def submit():
        # Get the selected client and order number
        nonlocal selected_item
        nonlocal order_number

        # Window Title "Please Select a Client"
        
        selected_item = listbox.get(listbox.curselection())
        order_number = order_entry.get()
        
        print("Selected Client:", selected_item)
        print("Order number:", order_number)

        # Close the window
        root.quit()
        root.destroy()

    root = tk.Tk()
    root.title("Order Form")
    root.geometry("300x300")

    # Create a listbox with the client names and a submit button
    label = tk.Label(root, text="Please Select a Client")


    listbox = tk.Listbox(root)
    for item in client_list:
        listbox.insert(tk.END, item)

    label.pack()
    listbox.pack()

    order_label = tk.Label(root, text="Order Number:")
    order_label.pack()

    order_entry = tk.Entry(root)
    order_entry.pack()

    submit_button = tk.Button(root, text="Submit", command=submit)
    submit_button.pack()

    root.mainloop()

    return selected_item, order_number

class CellLineSample():
    def __init__(self, sampleName, sampleDF, website, bestMatched):
        self.sampleName = sampleName
        self.sampleDF = sampleDF
        self.website = website
        self.bestMatched = bestMatched



def display_readme():
    try:
        with open('readme.txt', 'r') as file:
            readme_content = file.read()
            messagebox.showinfo("Read Me", readme_content)
    except FileNotFoundError:
        messagebox.showerror("Error", "readme.txt not found.")

    # Create the main Tkinter window
    root = tk.Tk()

    # Set window title and size
    root.title("My Program")
    root.geometry("300x200")

    # Create a button to show the Read Me message
    readme_button = tk.Button(root, text="Read Me", command=display_readme)
    readme_button.pack(pady=50)

    # Start the Tkinter event loop
    root.mainloop()

def show_done_window():
    def show_done_message():
        messagebox.showinfo("Done", "Cell Line ID script has finished running!")
        root.destroy()  # Close the "Done" window and exit the program

    # Create the main Tkinter window
    root = tk.Tk()

    # Set window title and size
    root.title("Done")
    root.geometry("300x200")

    # Create a label to display the "Done" message
    message_label = tk.Label(root, text="Process completed!")
    message_label.pack(pady=50)

    # Create a button to close the window and exit the program
    done_button = tk.Button(root, text="Done", command=show_done_message)
    done_button.pack()

    # Start the Tkinter event loop
    root.mainloop()

########################################################################################################################

if __name__ == "__main__":
    main()
