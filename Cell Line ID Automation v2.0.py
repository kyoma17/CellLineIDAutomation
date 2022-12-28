# Author: Kenny Ma
# Contact: 626-246-2233 or kyoma17@gmail.com
# Date: 2022-12-26
# Version: 2.0
# Description: This script will take in an excel file from the GMapper program and perform the ClimaSTR Cell Line ID script on each sample
# Requirements: Selenium, Pandas, BeautifulSoup, docx, tkinter, Firefox, geckodriver.exe
# there is a Requirements.txt file that can be used to install the required modules
# Use pip install -r Requirements.txt to install the required modules
# GekoDriver: Place geckodriver.exe in the same directory as this script
# Instructions: Run the script and select the input file from the GMapper program. The script will then perform the ClimaSTR Cell Line ID script on each sample and output a .docx file for each sample.
# The input file must have the following information: Sample Name, D5S818, D13S317, D7S820, D16S539, vWA, TH01, AMEL, TPOX, CSF1PO, D21S11


from tkinter import filedialog
import tkinter as tk
from tkinter import ttk
from bs4 import BeautifulSoup
import docx
import tkinter
import pandas
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.options import Options
import time
import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
# Import the webdriver and time modules


def main():
    print("Welcome to the ClimaSTR Cell Line ID script" '\n' "Please select the input file from the GMapper Program")
    # get file path from user using tkinter file dialog
    tkinter.Tk().withdraw()
    open file dialog in current directory and allow excel files only
    file_path = filedialog.askopenfilename(initialdir=".", title="Select file", filetypes=(
        ("Select CellLine ID file", "*.xlsx"), ("all files", "*.*")))

    # file_path = "C:/Users/kyo_m/Documents/Code/GP 10 input.xlsx"

    # read .xlsx into pandas dataframe
    df = pandas.read_excel(file_path)

    # Print dataframe and ask user to confirm correct file
    print(df)

    grouped_df = df.groupby("Sample Name")

    # for each group, perform the selenium script
    for each in grouped_df:
        sampleName = each[0]
        sampleDF = each[1]
        print("Processing " + sampleName + "...")

        clima_results = ClimaSTRSearch(sampleName, sampleDF)
        expasy_results = ExpasySTRSearch(sampleName, sampleDF)

        # Combine the results from ClimaSTR and ExpasySTR
        results = clima_results + expasy_results

        # Display the best matched samples to the user and ask for user input
        selectSample(results)


def selectSample(bestMatchedSamples):

    # unpack list of dictionaries
    Clima1st, Clima2nd, Clima3rd, Expasy1st, Expasy2nd, Expasy3rd = bestMatchedSamples

    # Display the best matched samples to the user and ask for user input
    window = tk.Tk()
    window.title("Select Result")

    tree = ttk.Treeview(window)

    columns = ("name", "_dataset", "_bMatchScore", "_bMatchName", "_bMatchCellLineNo", "D5S818_bM", "D13S317_bM",
            "D7S820_bM", "D16S539_bM", "vWA_bM", "TH01_bM", "AMEL_bM", "TPOX_bM", "CSF1PO_bM", "D21S11_bM")

    # Create the treeview
    tree = ttk.Treeview(window, columns=columns, show='headings')

    # Set the column widths
    for col in columns:
        tree.column(col, width=100, anchor=tk.W)

    # Set the column headings
    for col, heading in enumerate(columns):
        tree.heading(col, text=heading, anchor=tk.W)


    # Insert ClimaSTR results into treeview
    tree.insert("", 0, text="Clima 1st", values=("Clima1st", Clima1st["_dataset"], Clima1st["_bMatchScore"], Clima1st["_bMatchName"], Clima1st["_bMatchCellLineNo"], Clima1st["D5S818_bM"], Clima1st[
                "D13S317_bM"], Clima1st["D7S820_bM"], Clima1st["D16S539_bM"], Clima1st["vWA_bM"], Clima1st["TH01_bM"], Clima1st["AMEL_bM"], Clima1st["TPOX_bM"], Clima1st["CSF1PO_bM"], Clima1st["D21S11_bM"]))
    tree.insert("", 1, text="Clima 2nd", values=("Clima2nd", Clima2nd["_dataset"], Clima2nd["_bMatchScore"], Clima2nd["_bMatchName"], Clima2nd["_bMatchCellLineNo"], Clima2nd["D5S818_bM"], Clima2nd[
                "D13S317_bM"], Clima2nd["D7S820_bM"], Clima2nd["D16S539_bM"], Clima2nd["vWA_bM"], Clima2nd["TH01_bM"], Clima2nd["AMEL_bM"], Clima2nd["TPOX_bM"], Clima2nd["CSF1PO_bM"], Clima2nd["D21S11_bM"]))
    tree.insert("", 2, text="Clima3rd", values=("Clima3rd", Clima3rd["_dataset"], Clima3rd["_bMatchScore"], Clima3rd["_bMatchName"], Clima3rd["_bMatchCellLineNo"], Clima3rd["D5S818_bM"], Clima3rd[
                "D13S317_bM"], Clima3rd["D7S820_bM"], Clima3rd["D16S539_bM"], Clima3rd["vWA_bM"], Clima3rd["TH01_bM"], Clima3rd["AMEL_bM"], Clima3rd["TPOX_bM"], Clima3rd["CSF1PO_bM"], Clima3rd["D21S11_bM"]))

    # Insert ExpasySTR results into treeview
    tree.insert("", 3, text="Expasy 1st", values=("Expasy1st", Expasy1st["_dataset"], Expasy1st["_bMatchScore"], Expasy1st["_bMatchName"], Expasy1st["_bMatchCellLineNo"], Expasy1st["D5S818_bM"], Expasy1st[
                "D13S317_bM"], Expasy1st["D7S820_bM"], Expasy1st["D16S539_bM"], Expasy1st["vWA_bM"], Expasy1st["TH01_bM"], Expasy1st["AMEL_bM"], Expasy1st["TPOX_bM"], Expasy1st["CSF1PO_bM"], Expasy1st["D21S11_bM"]))
    tree.insert("", 4, text="Expasy 2nd", values=("Expasy2nd", Expasy2nd["_dataset"], Expasy2nd["_bMatchScore"], Expasy2nd["_bMatchName"], Expasy2nd["_bMatchCellLineNo"], Expasy2nd["D5S818_bM"], Expasy2nd[
                "D13S317_bM"], Expasy2nd["D7S820_bM"], Expasy2nd["D16S539_bM"], Expasy2nd["vWA_bM"], Expasy2nd["TH01_bM"], Expasy2nd["AMEL_bM"], Expasy2nd["TPOX_bM"], Expasy2nd["CSF1PO_bM"], Expasy2nd["D21S11_bM"]))
    tree.insert("", 5, text="Expasy 3rd", values=("Expasy3rd", Expasy3rd["_dataset"], Expasy3rd["_bMatchScore"], Expasy3rd["_bMatchName"], Expasy3rd["_bMatchCellLineNo"], Expasy3rd["D5S818_bM"], Expasy3rd[
                "D13S317_bM"], Expasy3rd["D7S820_bM"], Expasy3rd["D16S539_bM"], Expasy3rd["vWA_bM"], Expasy3rd["TH01_bM"], Expasy3rd["AMEL_bM"], Expasy3rd["TPOX_bM"], Expasy3rd["CSF1PO_bM"], Expasy3rd["D21S11_bM"]))

# Create a submit button, close window when clicked
    def submit():
        # Get the selected item
        item = tree.selection()[0]
        # Print the Name of the selected item
        print(tree.item(item, "values")[0])
        selection = tree.item(item, "values")[0]

        if selection == "Clima1st":
            fillTemplate(Clima1st)
        elif selection == "Clima2nd":
            fillTemplate(Clima2nd)
        elif selection == "Clima3rd":
            fillTemplate(Clima3rd)
        elif selection == "Expasy1st":
            fillTemplate(Expasy1st)
        elif selection == "Expasy2nd":
            fillTemplate(Expasy2nd)
        elif selection == "Expasy3rd":
            fillTemplate(Expasy3rd)
        else:
            print("Error: No selection made")
            
        window.destroy()
        window.quit()
        



    submit_button = tk.Button(window, text='Submit', command=submit)
    tree.pack()
    submit_button.pack()

    # Run the main loop
    window.mainloop()
    


def selectSamplez(bestMatchedSamples):
    # Displays a window with a list of best matched samples
    # User can select a sample and the sample name will be returned
    clima_1 = bestMatchedSamples[0]
    clima_2 = bestMatchedSamples[1]
    clima_3 = bestMatchedSamples[2]
    expasy_1 = bestMatchedSamples[3]
    expasy_2 = bestMatchedSamples[4]
    expasy_3 = bestMatchedSamples[5]

   # Create Dataframe
    df = pd.DataFrame(columns=["Dataset", "Score", "Name", "CellLineNo", "D5S818",
                      "D13S317", "D7S820", "D16S539", "vWA", "TH01", "AMEL", "TPOX", "CSF1PO", "D21S11"])

    # Add Clima 1st
    df.loc[len(df)] = [clima_1["_dataset"], clima_1["_bMatchScore"], clima_1["_bMatchName"], clima_1["_bMatchCellLineNo"], clima_1["D5S818_bM"], clima_1["D13S317_bM"],
                       clima_1["D7S820_bM"], clima_1["D16S539_bM"], clima_1["vWA_bM"], clima_1["TH01_bM"], clima_1["AMEL_bM"], clima_1["TPOX_bM"], clima_1["CSF1PO_bM"], clima_1["D21S11_bM"]]
    # Add Clima 2nd
    df.loc[len(df)] = [clima_2["_dataset"], clima_2["_bMatchScore"], clima_2["_bMatchName"], clima_2["_bMatchCellLineNo"], clima_2["D5S818_bM"], clima_2["D13S317_bM"],
                       clima_2["D7S820_bM"], clima_2["D16S539_bM"], clima_2["vWA_bM"], clima_2["TH01_bM"], clima_2["AMEL_bM"], clima_2["TPOX_bM"], clima_2["CSF1PO_bM"], clima_2["D21S11_bM"]]
    # Add Clima 3rd
    df.loc[len(df)] = [clima_3["_dataset"], clima_3["_bMatchScore"], clima_3["_bMatchName"], clima_3["_bMatchCellLineNo"], clima_3["D5S818_bM"], clima_3["D13S317_bM"],
                       clima_3["D7S820_bM"], clima_3["D16S539_bM"], clima_3["vWA_bM"], clima_3["TH01_bM"], clima_3["AMEL_bM"], clima_3["TPOX_bM"], clima_3["CSF1PO_bM"], clima_3["D21S11_bM"]]
    # Add Expasy 1st
    df.loc[len(df)] = [expasy_1["_dataset"], expasy_1["_bMatchScore"], expasy_1["_bMatchName"], expasy_1["_bMatchCellLineNo"], expasy_1["D5S818_bM"], expasy_1["D13S317_bM"],
                       expasy_1["D7S820_bM"], expasy_1["D16S539_bM"], expasy_1["vWA_bM"], expasy_1["TH01_bM"], expasy_1["AMEL_bM"], expasy_1["TPOX_bM"], expasy_1["CSF1PO_bM"], expasy_1["D21S11_bM"]]
    # Add Expasy 2nd
    df.loc[len(df)] = [expasy_2["_dataset"], expasy_2["_bMatchScore"], expasy_2["_bMatchName"], expasy_2["_bMatchCellLineNo"], expasy_2["D5S818_bM"], expasy_2["D13S317_bM"],
                       expasy_2["D7S820_bM"], expasy_2["D16S539_bM"], expasy_2["vWA_bM"], expasy_2["TH01_bM"], expasy_2["AMEL_bM"], expasy_2["TPOX_bM"], expasy_2["CSF1PO_bM"], expasy_2["D21S11_bM"]]
    # Add Expasy 3rd
    df.loc[len(df)] = [expasy_3["_dataset"], expasy_3["_bMatchScore"], expasy_3["_bMatchName"], expasy_3["_bMatchCellLineNo"], expasy_3["D5S818_bM"], expasy_3["D13S317_bM"],
                       expasy_3["D7S820_bM"], expasy_3["D16S539_bM"], expasy_3["vWA_bM"], expasy_3["TH01_bM"], expasy_3["AMEL_bM"], expasy_3["TPOX_bM"], expasy_3["CSF1PO_bM"], expasy_3["D21S11_bM"]]
    print(df)

    print("Clima 1st: " + clima_1["_dataset"] + " " + clima_1["_bMatchName"] + " " + clima_1["_bMatchCellLineNo"], clima_1["_bMatchScore"], clima_1["D5S818_bM"], clima_1["D13S317_bM"], clima_1["D7S820_bM"], clima_1["D16S539_bM"], clima_1["vWA_bM"], clima_1["TH01_bM"], clima_1["AMEL_bM"], clima_1["TPOX_bM"], clima_1["CSF1PO_bM"], clima_1["D21S11_bM"])
    print("Clima 2nd: " + clima_2["_dataset"] + " " + clima_2["_bMatchName"] + " " + clima_2["_bMatchCellLineNo"], clima_2["_bMatchScore"], clima_2["D5S818_bM"], clima_2["D13S317_bM"], clima_2["D7S820_bM"], clima_2["D16S539_bM"], clima_2["vWA_bM"], clima_2["TH01_bM"], clima_2["AMEL_bM"], clima_2["TPOX_bM"], clima_2["CSF1PO_bM"], clima_2["D21S11_bM"])
    print("Clima 3rd: " + clima_3["_dataset"] + " " + clima_3["_bMatchName"] + " " + clima_3["_bMatchCellLineNo"], clima_3["_bMatchScore"], clima_3["D5S818_bM"], clima_3["D13S317_bM"], clima_3["D7S820_bM"], clima_3["D16S539_bM"], clima_3["vWA_bM"], clima_3["TH01_bM"], clima_3["AMEL_bM"], clima_3["TPOX_bM"], clima_3["CSF1PO_bM"], clima_3["D21S11_bM"])
    print("Expasy 1st: " + expasy_1["_dataset"] + " " + expasy_1["_bMatchName"] + " " + expasy_1["_bMatchCellLineNo"], expasy_1["_bMatchScore"], expasy_1["D5S818_bM"], expasy_1["D13S317_bM"], expasy_1["D7S820_bM"], expasy_1["D16S539_bM"], expasy_1["vWA_bM"], expasy_1["TH01_bM"], expasy_1["AMEL_bM"], expasy_1["TPOX_bM"], expasy_1["CSF1PO_bM"], expasy_1["D21S11_bM"])
    print("Expasy 2nd: " + expasy_2["_dataset"] + " " + expasy_2["_bMatchName"] + " " + expasy_2["_bMatchCellLineNo"], expasy_2["_bMatchScore"], expasy_2["D5S818_bM"], expasy_2["D13S317_bM"], expasy_2["D7S820_bM"], expasy_2["D16S539_bM"], expasy_2["vWA_bM"], expasy_2["TH01_bM"], expasy_2["AMEL_bM"], expasy_2["TPOX_bM"], expasy_2["CSF1PO_bM"], expasy_2["D21S11_bM"])
    print("Expasy 3rd: " + expasy_3["_dataset"] + " " + expasy_3["_bMatchName"] + " " + expasy_3["_bMatchCellLineNo"], expasy_3["_bMatchScore"], expasy_3["D5S818_bM"], expasy_3["D13S317_bM"], expasy_3["D7S820_bM"], expasy_3["D16S539_bM"], expasy_3["vWA_bM"], expasy_3["TH01_bM"], expasy_3["AMEL_bM"], expasy_3["TPOX_bM"], expasy_3["CSF1PO_bM"], expasy_3["D21S11_bM"])

############################################################################################################


def ExpasySTRSearch(sampleName, sampleDF):
    # Takes in a pandas dataframe of a single sample and performs the selenium script against Expasy
    # Will return a pandas dataframe of the results of web scraping

    # Webpage input fields
    Amelogen = ["AMEL", "input-Amelogenin"]
    CSF1PO = ["CSF1PO", "input-CSF1PO"]
    D2S133 = ["D2S133", "input-D2S1338"]
    D3S135 = ["D3S1358", "input-D3S1358"]
    D5S818 = ["D5S818", "input-D5S818"]
    D7S820 = ["D7S820", "input-D7S820"]
    D8S1179 = ["D8S1179", "input-D8S1179"]
    D13S317 = ["D13S317", "input-D13S317"]
    D16S539 = ["D16S539", "input-D16S539"]
    D18S51 = ["D18S51", "input-D18S51"]
    D19S433 = ["D19S433", "input-D19S433"]
    D21S11 = ["D21S11", "input-D21S11"]
    FGA = ["FGA", "input-FGA"]
    PentaD = ["PentaD", "input-Penta_D"]
    PentaE = ["PentaE", "input-Penta_E"]
    TH01 = ["TH01", "input-TH01"]
    TPOX = ["TPOX", "input-TPOX"]
    vWA = ["vWA", "input-vWA"]

    master_list = [Amelogen, CSF1PO, D2S133, D3S135, D5S818, D7S820, D8S1179,
                   D13S317, D16S539, D18S51, D19S433, D21S11, FGA, PentaD, PentaE, TH01, TPOX, vWA]

    # Create a new instance of the Firefox driver
    options = Options()
    options.binary_location = r'C:\Program Files\Mozilla Firefox\firefox.exe'
    driver = webdriver.Firefox(
        executable_path=r'C:\WebDrivers\geckodriver.exe', options=options)

    # go to the Expasy STR website
    driver.get("https://www.cellosaurus.org/str-search/")
    time.sleep(1)

    # Input each allele into the webpage from the dataframe
    for each in master_list:
        marker = each[0]
        row = sampleDF.loc[sampleDF["Markers"] == marker]

        # combine Allele 1-4 from row into a single string, skip if NaN
        alleles = []
        for i in range(1, 5):
            allele = row[f"Allele{i}"].values
            if len(allele) > 0 and not pandas.isnull(allele[0]):
                alleles.append(str(allele[0]))

        allele = ",".join(alleles) if len(alleles) > 0 else ""

        # Input the allele into the webpage
        if len(allele) > 0:
            driver.find_element("id", each[1]).send_keys(allele)

    # Click the search button
    driver.find_element("id", "search").click()

    time.sleep(1)
    table = driver.find_element("id", "table-results")
    html = table.get_attribute("outerHTML")
    soup = BeautifulSoup(html, "html.parser")
    rows = soup.find_all("tr")
    header_row = rows[0]
    column_names = [th.text for th in header_row.find_all("th")]

    # Store the data in a list of lists
    data = []
    for row in rows:
        cells = row.find_all("td")
        data.append([cell.text for cell in cells])

    df = pd.DataFrame(data, columns=column_names)

    # Best match is row 3
    bestMatched_1 = df.iloc[2]
    bestMatched_2 = df.iloc[3]
    bestMatched_3 = df.iloc[4]

    driver.quit()

    replacementsDictionary1 = generateReplacementDictionary(
        sampleName, sampleDF, bestMatched_1, "Expasy")
    replacementsDictionary2 = generateReplacementDictionary(
        sampleName, sampleDF, bestMatched_2, "Expasy")
    replacementsDictionary3 = generateReplacementDictionary(
        sampleName, sampleDF, bestMatched_3, "Expasy")

    return [replacementsDictionary1, replacementsDictionary2, replacementsDictionary3]

    # # Fill the template with the data
    # fillTemplate(sampleName, replacementsDictionary)


def ClimaSTRSearch(sampleName, sampleDF):
    # Takes in a pandas dataframe of a single sample and performs the selenium script against Clima
    # Will return a pandas dataframe of the results of web scraping

    # Webpage input fields
    D5S818_list = ["D5S818", "D5S818_data1",
                   "D5S818_data2", "D5S818_data3", "D5S818_data4"]
    D13S317_list = ["D13S317", "D13S317_data1",
                    "D13S317_data2", "D13S317_data3", "D13S317_data4"]
    D7S820_list = ["D7S820", "D7S820_data1",
                   "D7S820_data2", "D7S820_data3", "D7S820_data4"]
    D16S539_list = ["D16S539", "D16S539_data1",
                    "D16S539_data2", "D16S539_data3", "D16S539_data4"]
    VWA_list = ["vWA", "VWA_data1", "VWA_data2", "VWA_data3", "VWA_data4"]
    TH01_list = ["TH01", "TH01_data1",
                 "TH01_data2", "TH01_data3", "TH01_data4"]
    Amelogenin_list = ["AMEL", "AMG_data1", "AMG_data2"]
    TPOX_list = ["TPOX", "TPOX_data1",
                 "TPOX_data2", "TPOX_data3", "TPOX_data4"]
    CSF1PO_list = ["CSF1PO", "CSF1PO_data1", "CSF1PO_data2",
                   "CSF1PO_data3", "CSF1PO_data4"]
    # D21S11_list = ["D21S11", "D21S11_data1", "D21S11_data2",
    #                "D21S11_data3", "D21S11_data4"]

    master_list = [D5S818_list, D13S317_list, D7S820_list, D16S539_list,
                   VWA_list, TH01_list, Amelogenin_list, TPOX_list, CSF1PO_list]

    # Create a new instance of the Firefox driver
    options = Options()
    options.binary_location = r'C:\Program Files\Mozilla Firefox\firefox.exe'
    driver = webdriver.Firefox(
        executable_path=r'C:\WebDrivers\geckodriver.exe', options=options)

    # Open a web browser and navigate to the website
    driver.get('http://bioinformatics.hsanmartino.it/clima2/index.php')

    # Wait for the page to load
    # time.sleep(1)

    # Input each allele into the webpage from the dataframe
    for each in master_list:
        marker = each[0]
        row = sampleDF.loc[sampleDF["Markers"] == marker]

        for i in range(1, len(each)):
            # skip nan values
            if pandas.isna(row["Allele" + str(i)].values[0]):
                continue

            allele = driver.find_element("id", each[i])

            allele.send_keys(row["Allele" + str(i)].values[0])

    # enter email and country and submit
    email = driver.find_element("id", "usr_email")
    email.send_keys("kyoma17@gmail.com")
    country = driver.find_element("id", "usr_country")
    country.send_keys("United States")

    # Find the submit button
    submit_button = driver.find_element(
        By.XPATH, "//input[@type='submit' and @value='submit']")

    # Click the submit button
    submit_button.click()

    # Wait for the page to load
    time.sleep(1)

    # retreive table from webpage and convert to pandas dataframe
    table = driver.find_element(By.XPATH, "(//table)[3]")
    html_table = table.get_attribute("outerHTML")
    soup = BeautifulSoup(html_table, "html.parser")

    # find the table
    souptable = soup.find("table")
    rows = souptable.find_all("tr")

    data = []
    for row in rows:
        cells = row.find_all("td")
        values = [cell.text for cell in cells]
        data.append(values)
    tableDF = pandas.DataFrame(data)

    PandasTableDF = pandas.read_html(html_table)[0]
    PandasTableDF.reset_index(drop=True, inplace=True)
    tableDF.columns = PandasTableDF.columns.get_level_values(0)

    # select the highest match
    bestMatched_1 = tableDF.iloc[2]
    bestMatched_2 = tableDF.iloc[3]
    bestMatched_3 = tableDF.iloc[4]

    # Close the browser
    driver.quit()

    # Generate the dictionary of replacements
    replacementsDictionary1 = generateReplacementDictionary(
        sampleName, sampleDF, bestMatched_1, "Clima")
    replacementsDictionary2 = generateReplacementDictionary(
        sampleName, sampleDF, bestMatched_2, "Clima")
    replacementsDictionary3 = generateReplacementDictionary(
        sampleName, sampleDF, bestMatched_3, "Clima")

    return [replacementsDictionary1, replacementsDictionary2, replacementsDictionary3]

    # # Transfer the data to the docx template
    # fillTemplate(sampleName, replacementsDictionary)


########################################################################################################################
def fillTemplate(replacementsDictionary):
    # Helper function to fill the template with the data from the dictionary
    sampleName = replacementsDictionary["_SAMPLE_NAME"]
    document = docx.Document('CellLineTemplateGP10.docx')
    # Iterate through the tables in the document and replace the text from the dictionary
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in replacementsDictionary.keys():
                        if key in paragraph.text:
                            # Enter the replacement text in the document, if the value is None or nan, then replace with an empty string
                            if replacementsDictionary[key] is None or pandas.isna(replacementsDictionary[key]):
                                paragraph.text = paragraph.text.replace(
                                    key, "")
                            else:
                                paragraph.text = paragraph.text.replace(
                                    key, str(replacementsDictionary[key]))

    # Save the modified document
    document.save(sampleName + "_" +
                  replacementsDictionary["website"] + ".docx")
    print("Done with " + replacementsDictionary["website"] + " " + sampleName)


def generateReplacementDictionary(sampleName, sampleDF, bestMatched, website):
    # Generate Replacement Dictionary for the Template
    # Clima or Expasy dictionary for the template
    if website == "Clima":
        replacementsDictionary = {
            # Main Info
            "_SAMPLE_NAME": sampleName,
            "website": "Clima2",


            # Data from the Results highest scoring match
            "_dataset": bestMatched["Dataset"],
            "_bMatchScore": bestMatched["% Match"],
            "_bMatchName": bestMatched["Name"],
            "_bMatchCellLineNo": bestMatched["Cat. No."],

            "D5S818_bM": bestMatched["D5S818"],
            "D13S317_bM": bestMatched["D13S317"],
            "D7S820_bM": bestMatched["D7S820"],
            "D16S539_bM": bestMatched["D16S539"],
            "vWA_bM": bestMatched["VWA"],
            "TH01_bM": bestMatched["TH01"],
            "AMEL_bM": bestMatched["AMG"],
            "TPOX_bM": bestMatched["TPOX"],
            "CSF1PO_bM": bestMatched["CSF1PO"],
            "D21S11_bM": bestMatched["D21S11"],
        }

    elif website == "Expasy":
        replacementsDictionary = {

            "_SAMPLE_NAME": sampleName,
            "website": "Expasy",

            # Data from the Results highest scoring match
            "_dataset": "Expasy",
            "_bMatchScore": bestMatched["Score"],
            "_bMatchName": bestMatched["Accession"],
            "_bMatchCellLineNo": bestMatched["Name"],

            "D5S818_bM": bestMatched["D5S818"],
            "D13S317_bM": bestMatched["D13S317"],
            "D7S820_bM": bestMatched["D7S820"],
            "D16S539_bM": bestMatched["D16S539"],
            "vWA_bM": bestMatched["vWA"],
            "TH01_bM": bestMatched["TH01"],
            "AMEL_bM": bestMatched["Amel"],
            "TPOX_bM": bestMatched["TPOX"],
            "CSF1PO_bM": bestMatched["CSF1PO"],
            "D21S11_bM": bestMatched["D21S11"],
        }

    # Replacement Dictionary for the Template from  Data Input
    markers = ["D5S818", "D13S317", "D7S820", "D16S539",
               "vWA", "TH01", "AMEL", "TPOX", "CSF1PO", "D21S11"]
    allele_columns = ["Allele1", "Allele2", "Allele3", "Allele4"]

    template_dict = {}

    for marker in markers:
        for i, allele_column in enumerate(allele_columns, start=1):
            key = f"{marker}_{i}"
            value = sampleDF.loc[sampleDF["Markers"]
                                 == marker][allele_column].values[0]
            template_dict[key] = value

    replacementsDictionary.update(template_dict)

    return replacementsDictionary


main()
