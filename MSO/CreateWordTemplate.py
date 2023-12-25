from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize a Document object
doc = Document()

# Add a heading
heading = doc.add_heading(level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run('Database Sample Match Report')
run.bold = True

# Function to add a row of data to the table
def add_row(table, data):
    row_cells = table.add_row().cells
    for index, cell in enumerate(row_cells):
        cell.text = str(data[index])

# Function to create a formatted table with headers and sample data
def create_table(doc, headers, num_columns, num_rows):
    table = doc.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for _ in range(num_rows):
        add_row(table, [''] * num_columns)  # Add empty rows for data to be filled in
    return table

# Add table for sample identification
sample_id_headers = ['Sample ID:', '_SAMPLE_NAME', 'Dataset Best Match Cell Name:', '_bMatchName']
sample_id_table = create_table(doc, sample_id_headers, 4, 1)

# Add table for best match scores
score_headers = ['Database Best Match Score:', '_bMatchScore', 'Dataset Best Match Cell Line No:', '_bMatchCellLineNo']
score_table = create_table(doc, score_headers, 4, 1)

# Add Methodology text
doc.add_paragraph('Authenticity for all cell lines used in this study was performed using the Promega GenePrint 10 '
                  'Methodology: System and following the protocol described in ANSI/ATCC ASN-0002-2011. The STR alleles '
                  'were searched on the _dataset Database.')

# Add Marker Characterization table
marker_headers = ['Marker', 'Allele 1', 'Allele 2', 'Allele 3', 'Allele 4', 'Dataset Best Match Profile']
marker_table = create_table(doc, marker_headers, 6, 8)  # Adjust the number of rows as needed

# Save the document
print("Saving document...")
doc.save('tesdfdft.docx')


