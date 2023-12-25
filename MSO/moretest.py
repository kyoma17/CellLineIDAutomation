
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize a Document object
doc = Document()

# Function to set a cell's background color
def set_cell_background(cell, fill):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Add a table with a specific layout as shown in the user's image
table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'

# Define cell text and background colors
cell_text = [
    ("sampleNumber", "Sample ID:", "_SAMPLE_NAME", "_dataset Best Match Cell Name:"),
    ("Database Best Match Score:", "_bMatchScore", "_dataset Best Match Cell Line No:", "_bMatchCellLineNo")
]
background_colors = ["D3D3D3", "FFFFFF", "FF0000", "FFFFFF"]  # Light grey, white, and red for simplicity

# Populate the table with the specified text and background colors
for row_index, row in enumerate(table.rows):
    for col_index, cell in enumerate(row.cells):
        # Set cell text
        cell.text = cell_text[row_index][col_index]
        
        # Set background color
        set_cell_background(cell, background_colors[col_index])

        # Center align text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set font to bold for the placeholders
        if "_" in cell.text:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True

# Save the document
doc.save('test.docx')

